from typing import List

import tqdm
"""
tqdm 用于显示进度条的python库，可以用在循环，迭代，文件读写等操作中显示进度，
"""
from langchain_community.document_loaders.unstructured import UnstructuredFileLoader


class RapidOCRDocLoader(UnstructuredFileLoader):
    """doc文档加载器扩展，可识别图片中的文字信息

    
    """
    def _get_elements(self) -> List:
        def doc2text(filepath):
            from io import BytesIO

            import numpy as np
      
            from docx import Document, ImagePart
            """
            docx模块是用于处理word文档的python库，
                Document类是docx模块中的类，用于读取和操作Word文档。通过使用Document类，可以读取Word文档的内容、样式、表格、图片等元素，并进行相应操作
                ImagePart类是docx模块中的一个内部类，用于表示word文档中的图片元素，通过使用ImagePart类，可以获取图片的路径，尺寸，位置等信息，并对图片进行操作
            """
            from docx.oxml.table import CT_Tbl
            """
            CT_Tbl用来表示Word文档中的表格元素。通过使用CT_Tbl类可以获取表格的行、列、单元格等信息
            """
            from docx.oxml.text.paragraph import CT_P
            """
            CT_P用于表示Word中段落元素，通过CT_P类，可以获取段落的样式，内容，缩进，对齐等信息，并对段落进行操作  
            """
            from docx.table import Table, _Cell
            from docx.text.paragraph import Paragraph
            from PIL import Image
            """
            导入Pillow库中的Image模块，提供了对各种图像格式的读取、写入、转换、滤镜等操作的支持
            """
            from rapidocr_onnxruntime import RapidOCR
            """
            快速 OCR 处理，OCR原理：
                输入原图 -> 文本检测模型 -> 得到多个文本区域框 -> 将每个框剪裁并修正 -> 分别送入文本识别模型 -> 得到所有识别结果 -> 输出结构化wenb
            """
            # 、、初始化OCR模型实例
            ocr = RapidOCR()
            # 、、获取文档
            doc = Document(filepath)
            resp = ""

            # 、、 生成一个迭代器，这个迭代器会以段落paragraph或者table的维度来生成
            def iter_block_items(parent):
                from docx.document import Document

                if isinstance(parent, Document):
                    parent_elm = parent.element.body
                elif isinstance(parent, _Cell):
                    # 、、这个分支应该不会被触发走进来，我没有看到关于iter_block_items的循环调用的地方，只调用了一次，并且只传入了doc文档，应该会直走上面的分支
                    parent_elm = parent._tc
                else:
                    raise ValueError("RapidOCRDocLoader parse fail")

                for child in parent_elm.iterchildren():
                    if isinstance(child, CT_P):
                        yield Paragraph(child, parent)
                    elif isinstance(child, CT_Tbl):
                        yield Table(child, parent)

            b_unit = tqdm.tqdm(
                total=len(doc.paragraphs) + len(doc.tables),
                desc="RapidOCRDocLoader block index: 0",
            )
            """
            for i in tqdm.tqdm(rang(100)) 将执行循环操作，显示以一个进度条，显示循环的进度；
                2个可选参数：desc(描述信息)和ncols(进度条的列数)
                  1、interable: 迭代对象，例如列表、字典、字符串等
                  2、desc: 进度条的描述信息
            返回的是一个可迭代对象 b_unit
            """
            # 、、循环 iter_block_items(doc)迭代器（一个以段落和table组成的可迭代对象，并且内部用yield来控制，确保边用边加载生成可迭代对象，如table，paragraph）
            for i, block in enumerate(iter_block_items(doc)):
                # 、、进度条
                b_unit.set_description("RapidOCRDocLoader  block index: {}".format(i))
                b_unit.refresh()
                # 、、如果是段落
                if isinstance(block, Paragraph):
                    # 、、1、进行简单的文本清洗，去除前后空格后累加到resp中，并且加入换行符（partition_text 会识别换行符作为潜在的分割点）
                    resp += block.text.strip() + "\n"
                    # 、、2、获取段落中所有的图片
                    images = block._element.xpath(".//pic:pic")  # 获取所有图片
                    # 、、3、循环图片
                    for image in images:
                        for img_id in image.xpath(".//a:blip/@r:embed"):  # 获取图片id
                            part = doc.part.related_parts[
                                img_id
                            ]  # 、、根据图片id获取对应的图片
                            if isinstance(part, ImagePart):
                                image = Image.open(BytesIO(part._blob))
                                # 、、使用ocr实例获取ocr对图片的识别结果，（图片中的文字，具体看模型，一定有模型可以识别出文字和图像比如苹果，并且同时转化为文字，输出文字进行存储，上一个项目的多模态识别就可以做到）
                                result, _ = ocr(np.array(image))
                                if result:
                                    ocr_result = [line[1] for line in result]
                                    # 、、对OCR的识别结果进行存储累加到resp中
                                    resp += "\n".join(ocr_result)
                # 、、对表格信息进行提取，并累加到resp中
                elif isinstance(block, Table):
                    for row in block.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                resp += paragraph.text.strip() + "\n"
                # 、、 进度条更新 b_unit.update(1)表示将进度条的进度增加1，也就是说，我们的进度条是以段落为单位，每加载一个段落，就更新一次进度条
                # 、、update()方法只能增加进度，不会将进度条的进度减少。如果要将进度条的进度减少，可以使用total参数来设置进度条的总进度，然后使用update()方法来更新进度条的进度
                b_unit.update(1)
            # 将所有收集到的信息返回  （包括段落文字，段落文字中的图像中的文字，表格中的文字）
            return resp

        text = doc2text(self.file_path)
        from unstructured.partition.text import partition_text
        # 、、将连续的文本流智能地分割成有意义的文档块（chunks），这些块在RAG系统中更适合进行向量化和检索。
        # 、、最终返回的是List[document]
        #    每个document包括：
        #       text：分割后的文本，
        #       metadata:元数据（如来源，位置等）
        #       其它结构化信息
        # 这样处理后的文档列表可以直接用于：
        #   文本嵌入（embedding）
        #   向量数据库存储
        #   相似性检索
        #   上下文构建
        return partition_text(text=text, **self.unstructured_kwargs)


if __name__ == "__main__":
    loader = RapidOCRDocLoader(file_path="../tests/samples/ocr_test.docx")
    # 、、这里测试加载并且打印到控制台中，
    # 、、load方法继承自UnstructuredFileLoader <- UnstructuredBaseLoader（这个里面实现了具体的lazy_load方法，调用load会触发调用链式调用 lazy_load方法） <- BaseLoader (BaseLoader这个抽象类定义了load方法)
    # 、、在langchain中load方法将lazy_load放到了一个List中并执行，
    # 、、lazy_load的执行是利用yeild来实现懒加载上面partition_text生成的结构化数据，以实现内存友好的效果，将每个文本元素document都  成为独立的Document，依次执行直到所有的结构化数据执行完毕，输出
    docs = loader.load()
    # docs为一个List[Document]，每个Document包括： page_content： 文本内存，字符串类型，metadata： 元数据，字典类型，包含文件名、来源，位置等，
    # Document(
    #     page_content="具体的文本内容",  # 字符串
    #     metadata={                     # 字典，包含各种元数据
    #         'source': '文件路径',
    #         'filename': '文件名',
    #         'filetype': 'text',
    #         'languages': ['zh'],
    #         'category': '文本类型',
    #         'page_number': 1,
    #         # ... 其他元数据
    #     }
    # )
    print(docs)
